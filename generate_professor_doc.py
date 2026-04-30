from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def h(text, level=1):
    return doc.add_heading(text, level=level)

def sec(label, text):
    p = doc.add_paragraph()
    p.add_run(label + '  ').bold = True
    p.add_run(text)

def b(text):
    return doc.add_paragraph(text, style='List Bullet')

# Title
t = doc.add_heading('Thesis Outline', 0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph('Working title: A Digital-Twin Simulation Framework for Joint Inventory and Transport Optimisation in Multi-Echelon Supply Chains')
p.runs[0].italic = True
doc.add_paragraph('Author: Sakthe Balan  |  Programme: Dual Degree Project (DDP)')
doc.add_paragraph('')

# Abstract
h('Abstract')
doc.add_paragraph('(To be written last) — gap between closed-form MEIO models and real supply chains; framework built, experiments run, and key results.')

# Ch 1
h('Chapter 1 — Introduction')
sec('1.1', 'Problem: why multi-echelon inventory is harder than single-echelon.')
sec('1.2', 'Motivation for a digital twin over closed-form models.')
sec('1.3', 'Contributions: 7-policy framework, joint optimiser, Pareto analysis, disruption modelling.')
sec('1.4', 'Thesis structure.')
sec('1.5', 'Development trajectory: prototype → mid-sem validation → 5-SKU scale-up → joint optimisation.')

# Ch 2
h('Chapter 2 — Literature Review')
sec('2.1', 'Single-echelon foundations: newsvendor, base-stock, (s,S), EOQ.')
sec('2.2', 'Multi-echelon optimality: Clark & Scarf (1960), echelon base-stock.')
sec('2.3', 'Periodic-review and cyclic policies: (R,S), (R,Q), (k,m)-cycle.')
sec('2.4', 'Bullwhip effect: causes and CV²-ratio metric.')
sec('2.5', 'Transport consolidation in MEIO models.')
sec('2.6', 'Simulation-optimisation: why Optuna TPE suits mixed-integer spaces.')
sec('2.7', 'Existing tools and the gap this DDP fills.')

# Ch 3
h('Chapter 3 — Digital-Twin Architecture')
sec('3.1', 'Network topology: 1-supplier → 1-warehouse → 3-retailers (1N3); arbitrary DAG support.')
sec('3.2', 'Simulation loop: daily tick — receive → demand → backlog → policy → order → dispatch.')
sec('3.3', 'Node state: on-hand, pipeline, external/internal backlog.')
sec('3.4', 'Transport layer: GreedyLoadPlanner with volume + weight caps, min utilisation threshold.')
sec('3.5', 'Cost model: holding + transport + ordering + shortage, warm-up-excluded.')
sec('3.6', 'Input methods:')
b('JSON config — hand-authored scenario files, validated by config_validator.')
b('CSV pipeline — 6 spreadsheet files (sim_config, products, nodes, policies, routes, demand_config) compiled into JSON by a builder script.')
b('Demand data — per-retailer daily CSVs derived from M5 Walmart dataset; synthetic demand_shock.csv for validation.')
sec('3.7', 'Supply disruption modelling: configurable node downtime, backlog accumulation.')
sec('3.8', 'Iterative evolution: validated on simple 2W×3R single-SKU prototype before scaling to 1N3/5-SKU.')

# Ch 4
h('Chapter 4 — Inventory Policies')
doc.add_paragraph('All 7 policies share a common API: order_qty(on_hand, backlog, pipeline, t) → int.')
sec('4.1', 'Base-stock: newsvendor analytical baseline.')
sec('4.2', '(s,S) and (R,Q): continuous-review reorder-point policies.')
sec('4.3', 'Periodic-review (R,S) and order-up-to with phase offset.')
sec('4.4', '(k,m)-cycle: k orders per m-day cycle.')
sec('4.5', 'Echelon base-stock (Clark & Scarf): DFS over downstream nodes for echelon inventory position.')
sec('4.6', 'Discussion: trade-offs across the policy zoo.')

# Ch 5
h('Chapter 5 — Optimisation Framework')
sec('5.1', 'Optuna TPE: sample-efficient on mixed-integer spaces.')
sec('5.2', 'Inventory search space: per-(node, SKU) base-stock levels.')
sec('5.3', 'Transport search space: per-lane min dispatch utilisation.')
sec('5.4', 'Joint search: single study over combined inventory + transport space.')
sec('5.5', 'Fill-rate constraint: soft penalty; best feasible trial selected.')
sec('5.6', 'Pareto exploration: re-run at fill-rate targets 80%–98%.')
sec('5.7', 'Re-evaluation discipline: best params re-run in fresh simulator to guard against overfitting.')

# Ch 6
h('Chapter 6 — Experimental Methodology')
sec('6.0', 'Preliminary validation (mid-semester): 4 experiments on 2W×3R single-SKU topology.')
b('EV1 — Constant-demand: validates ordering logic.')
b('EV2 — Demand shock: validates pipeline lead-time dynamics.')
b('EV3 — Policy comparison with/without transport capacity.')
b('EV4 — Capacity scaling on M5 data: non-linear cost–service knee.')
sec('6.1', 'Dataset: M5 Walmart demand (5 SKUs, 365 days).')
sec('6.2', 'Warm-up exclusion, multi-seed CIs, bullwhip metric.')
sec('6.3', 'Main experiments E0–E4:')
tbl = doc.add_table(rows=1, cols=3)
tbl.style = 'Table Grid'
hdr = tbl.rows[0].cells
hdr[0].text = 'ID'; hdr[1].text = 'Title'; hdr[2].text = 'Hypothesis'
for row in [
    ('E0',  'Analytical baseline',          'High transport cost, acceptable fill'),
    ('E1a', 'Fixed dispatch threshold',     'Cuts transport cost, hurts fill rate'),
    ('E1b', 'Transport-only optimisation',  'Limited by fixed inventory parameters'),
    ('E2',  'Inventory-only optimisation',  'Better fill, modest cost improvement'),
    ('E3',  'Joint optimisation',           'Lowest cost at ≥92% fill'),
    ('E4',  'Pareto frontier',              'Diminishing returns near 100% fill'),
]:
    r = tbl.add_row().cells
    r[0].text = row[0]; r[1].text = row[1]; r[2].text = row[2]
doc.add_paragraph('')

# Ch 7
h('Chapter 7 — Results')
sec('7.0', 'Preliminary validation results (EV1–EV4):')
b('EV1: inventory converges to target within L+1 days.')
b('EV2: upstream spikes delayed 1–3 days, amplitude amplified.')
b('EV3: capacity constraints cause (s,S) to outperform base-stock on transport cost.')
b('EV4: 0.25× → 49.3% fill/$73k; 1× → 98%/$184k; 4× → 99.3%/$712k — identifies cost-service knee.')
sec('7.1–7.4', 'E0–E2 results: transport dominates cost; single-axis optimisation has clear limits.')
sec('7.5', 'E3 (headline): joint optimisation reduces cost vs E0 at ≥92% fill.')
sec('7.6', 'E4: Pareto frontier — marginal cost per service-level point.')
sec('7.7', 'Bullwhip analysis: ratios per experiment; partial suppression under joint optimisation.')
sec('7.8', 'Statistical robustness: multi-seed CIs; E3 improvement is significant.')
sec('7.9', 'Disruption sensitivity (planned): 14-day supplier outage impact on E0 vs E3.')

# Ch 8
h('Chapter 8 — Discussion')
sec('8.1', 'Research questions answered: joint beats single-axis; disruption handled better with elevated DC stock; bullwhip partially suppressed.')
sec('8.2', 'Practical recommendations: co-optimise inventory and dispatch; use Pareto frontier for SLA negotiation.')
sec('8.3', 'Limitations: deterministic lead times, no transshipment, offline demand.')

# Ch 9
h('Chapter 9 — Conclusion and Future Work')
sec('9.1', 'Summary of contributions with quantitative results.')
sec('9.2', 'Future work: stochastic lead times, forecasting layer, RL policies, lateral transshipment.')
sec('9.3', 'Closing remarks.')

# Appendices
h('Appendices')
b('A — Configuration schema reference.')
b('B — Per-experiment policy parameters.')
b('C — Test suite summary (~80 pytest tests).')
b('D — CLI commands to reproduce results.')

doc.save('/home/sakthe/Desktop/DDP/thesis_outline_for_professor.docx')
print('Saved.')
